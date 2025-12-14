from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, send_from_directory, jsonify
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import os
import sys
import io
import threading
import tempfile
import random
import json
import time
from datetime import datetime

# 创建Flask应用
app = Flask(__name__)
app.secret_key = 'your_secret_key'  # 用于会话管理的密钥

# 设置字符编码
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = 'application/json;charset=utf-8'

# 确保所有响应使用UTF-8编码
@app.after_request
def set_encoding(response):
    response.headers['Content-Type'] = 'text/html; charset=utf-8'
    return response

# 添加自定义过滤器：fromjson
@app.template_filter('fromjson')
def fromjson_filter(s):
    import json
    if s and s != 'null':
        try:
            return json.loads(s)
        except:
            return {}
    return {}

# 配置数据库
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///user_management.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# 初始化数据库
db = SQLAlchemy(app)

# 权限定义
PERMISSIONS = {
    'student': {
        'name': '学生',
        'permissions': [
            '学习课程',
            '提交作业',
            '查看成绩',
            '参与讨论',
            'AI助教答疑',
            '项目组管理'
        ]
    },
    'teacher': {
        'name': '教师',
        'permissions': [
            '创建课程',
            '编辑课程',
            '批改作业',
            '管理学生',
            '发布成绩',
            '参与讨论',
            'AI测验管理',
            'AI助教答疑',
            '项目组管理'
        ]
    },
    'ai-assistant': {
        'name': 'AI 助教',
        'permissions': [
            '回答问题',
            '批改测验',
            '提供学习建议',
            '辅助管理课程'
        ]
    }
}

# 用户模型
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), nullable=False)
    student_id = db.Column(db.String(20), unique=True, nullable=False)  # 学号/工号
    
    # 统计信息（简化版）
    courses = db.Column(db.Integer, default=0)
    assignments = db.Column(db.Integer, default=0)
    completed_assignments = db.Column(db.Integer, default=0)
    average_score = db.Column(db.Float, default=0.0)
    student_count = db.Column(db.Integer, default=0)
    graded_assignments = db.Column(db.Integer, default=0)
    questions_answered = db.Column(db.Integer, default=0)
    quizzes_graded = db.Column(db.Integer, default=0)
    suggestions_provided = db.Column(db.Integer, default=0)

# 课程模型
class Course(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    course_code = db.Column(db.String(20), unique=True, nullable=False)  # 课程代码
    title = db.Column(db.String(100), nullable=False)  # 课程标题
    description = db.Column(db.Text, nullable=True)  # 课程描述
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)  # 授课老师ID
    credit = db.Column(db.Float, default=2.0)  # 学分
    created_at = db.Column(db.DateTime, default=datetime.utcnow)  # 创建时间
    video_path = db.Column(db.String(255), nullable=True)  # 视频文件路径
    
    # 关系
    teacher = db.relationship('User', backref=db.backref('courses_taught', lazy=True))

# 学生选课关联表
class StudentCourse(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)  # 学生ID
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'), nullable=False)  # 课程ID
    enrolled_at = db.Column(db.DateTime, default=datetime.utcnow)  # 选课时间
    
    # 关系
    student = db.relationship('User', backref=db.backref('enrolled_courses', lazy=True, overlaps="courses"))
    course = db.relationship('Course', backref=db.backref('students', lazy=True))

# 题目模型
class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    qid = db.Column(db.String(50), unique=True, nullable=False)
    knowledge_point = db.Column(db.String(100), nullable=False)
    difficulty = db.Column(db.String(20), nullable=False)
    qtype = db.Column(db.String(20), nullable=False)  # single, multiple, judge, short
    content = db.Column(db.Text, nullable=False)
    answer = db.Column(db.Text, nullable=False)
    options = db.Column(db.Text, nullable=True)  # JSON格式存储选项
    score_std = db.Column(db.Text, nullable=True)  # 评分标准，用于主观题
    source = db.Column(db.String(20), default="ai_generated")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# 测验模型
class Quiz(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    quiz_id = db.Column(db.String(50), unique=True, nullable=False)
    title = db.Column(db.String(100), nullable=False)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'), nullable=False)
    knowledge_points = db.Column(db.Text, nullable=False)  # JSON格式存储知识点范围
    difficulty = db.Column(db.String(20), nullable=False)
    time_limit = db.Column(db.Integer, nullable=False)  # 时间限制（分钟）
    anti_cheat = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # 关系
    teacher = db.relationship('User', backref=db.backref('quizzes', lazy=True))
    course = db.relationship('Course', backref=db.backref('quizzes', lazy=True))

# 测验题目关联表
class QuizQuestion(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    quiz_id = db.Column(db.Integer, db.ForeignKey('quiz.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    score = db.Column(db.Float, nullable=False)  # 本题分值
    
    # 关系
    quiz = db.relationship('Quiz', backref=db.backref('quiz_questions', lazy=True))
    question = db.relationship('Question', backref=db.backref('quiz_questions', lazy=True))

# 学生测验记录模型
class StudentQuiz(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    quiz_id = db.Column(db.Integer, db.ForeignKey('quiz.id'), nullable=False)
    start_time = db.Column(db.DateTime, nullable=False)
    end_time = db.Column(db.DateTime, nullable=True)
    total_score = db.Column(db.Float, default=0.0)
    status = db.Column(db.String(20), default="in_progress")  # in_progress, completed, submitted
    cheat_count = db.Column(db.Integer, default=0)
    is_timeout = db.Column(db.Boolean, default=False)
    
    # 关系
    student = db.relationship('User', backref=db.backref('student_quizzes', lazy=True))
    quiz = db.relationship('Quiz', backref=db.backref('student_quizzes', lazy=True))

# 学生答题记录模型
class StudentAnswer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_quiz_id = db.Column(db.Integer, db.ForeignKey('student_quiz.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    student_answer = db.Column(db.Text, nullable=False)
    is_correct = db.Column(db.Boolean, nullable=True)
    score = db.Column(db.Float, default=0.0)
    spend_time = db.Column(db.Float, default=0.0)  # 答题用时（秒）
    submit_time = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    student_quiz = db.relationship('StudentQuiz', backref=db.backref('student_answers', lazy=True))
    question = db.relationship('Question', backref=db.backref('student_answers', lazy=True))

# 错题库模型
class ErrorQuestionBank(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    first_error_time = db.Column(db.DateTime, default=datetime.utcnow)
    last_error_time = db.Column(db.DateTime, default=datetime.utcnow)
    error_count = db.Column(db.Integer, default=1)
    notes = db.Column(db.Text, nullable=True)
    
    # 关系
    student = db.relationship('User', backref=db.backref('error_questions', lazy=True))
    question = db.relationship('Question', backref=db.backref('error_questions', lazy=True))

# 消息模型
class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sender_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    receiver_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    is_read = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    sender = db.relationship('User', foreign_keys=[sender_id], backref=db.backref('sent_messages', lazy=True))
    receiver = db.relationship('User', foreign_keys=[receiver_id], backref=db.backref('received_messages', lazy=True))

# 项目组模型
class ProjectGroup(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    creator_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    creator = db.relationship('User', backref=db.backref('created_projects', lazy=True))

# 项目组成员关联表
class GroupMember(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    group_id = db.Column(db.Integer, db.ForeignKey('project_group.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    role = db.Column(db.String(20), default='member')  # member, admin, teacher
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    group = db.relationship('ProjectGroup', backref=db.backref('members', lazy=True))
    user = db.relationship('User', backref=db.backref('project_groups', lazy=True))

# 项目组讨论消息模型
class GroupDiscussion(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    group_id = db.Column(db.Integer, db.ForeignKey('project_group.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    group = db.relationship('ProjectGroup', backref=db.backref('discussions', lazy=True))
    user = db.relationship('User', backref=db.backref('group_discussions', lazy=True))

# AI辅导会话模型
class AITutoringSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'), nullable=True)
    session_type = db.Column(db.String(50), nullable=False)  # course, assignment, general
    topic = db.Column(db.String(200), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    user = db.relationship('User', backref=db.backref('ai_sessions', lazy=True))
    course = db.relationship('Course', backref=db.backref('ai_sessions', lazy=True))

# AI辅导消息模型
class AITutoringMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    session_id = db.Column(db.Integer, db.ForeignKey('ai_tutoring_session.id'), nullable=False)
    sender = db.Column(db.String(20), nullable=False)  # user, ai
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # 关系
    session = db.relationship('AITutoringSession', backref=db.backref('messages', lazy=True))

# 创建数据库表
with app.app_context():
    db.create_all()
    
    # 创建默认老师用户
    if not User.query.filter_by(username='default_teacher').first():
        default_teacher = User(
            username='default_teacher',
            password=generate_password_hash('123456', method='pbkdf2:sha256'),
            role='teacher',
            student_id='T001'
        )
        db.session.add(default_teacher)
        db.session.commit()
    else:
        default_teacher = User.query.filter_by(username='default_teacher').first()
    
    # 添加预设课程
    preset_courses = [
        {
            'course_code': 'ART001',
            'title': '造型艺术',
            'description': '本课程介绍造型艺术的基本概念、历史发展和创作方法，培养学生的艺术审美和创作能力。',
            'credit': 2.0
        },
        {
            'course_code': 'DES001',
            'title': '设计思维',
            'description': '本课程教授设计思维的核心方法和实践技巧，培养学生的创新思维和问题解决能力。',
            'credit': 3.0
        },
        {
            'course_code': 'DES002',
            'title': '设计基础',
            'description': '本课程介绍设计的基本原理、构成要素和表现技法，为学生打下坚实的设计基础。',
            'credit': 2.0
        },
        {
            'course_code': 'ART002',
            'title': '现代设计史',
            'description': '本课程讲述现代设计的发展历程、重要流派和代表人物，帮助学生理解设计的历史脉络和演变规律。',
            'credit': 3.0
        }
    ]
    
    for course_data in preset_courses:
        if not Course.query.filter_by(course_code=course_data['course_code']).first():
            new_course = Course(
                course_code=course_data['course_code'],
                title=course_data['title'],
                description=course_data['description'],
                teacher_id=default_teacher.id,
                credit=course_data['credit']
            )
            db.session.add(new_course)
    db.session.commit()

# 登录页面
@app.route('/')
def login():
    return render_template('login.html')

# 登录处理
@app.route('/login', methods=['POST'])
def login_process():
    username = request.form['username']
    password = request.form['password']
    
    user = User.query.filter_by(username=username).first()
    
    if user and check_password_hash(user.password, password):
        session['user_id'] = user.id
        session['username'] = user.username
        session['role'] = user.role
        return redirect(url_for('dashboard'))
    else:
        flash('用户名或密码错误')
        return redirect(url_for('login'))

# 注册页面
@app.route('/register')
def register():
    return render_template('register.html')

# 注册处理
@app.route('/register', methods=['POST'])
def register_process():
    username = request.form['username']
    password = request.form['password']
    confirm_password = request.form['confirm_password']
    role = request.form['role']
    student_id = request.form['student_id']
    
    # 验证密码
    if password != confirm_password:
        flash('两次输入的密码不一致')
        return redirect(url_for('register'))
    
    # 检查用户名是否已存在
    if User.query.filter_by(username=username).first():
        flash('用户名已存在')
        return redirect(url_for('register'))
    
    # 检查学号/工号是否已存在
    if User.query.filter_by(student_id=student_id).first():
        flash('学号/工号已被使用')
        return redirect(url_for('register'))
    
    # 创建新用户
    hashed_password = generate_password_hash(password, method='pbkdf2:sha256')
    new_user = User(
        username=username,
        password=hashed_password,
        role=role,
        student_id=student_id
    )
    
    db.session.add(new_user)
    db.session.commit()
    
    flash('注册成功，请登录')
    return redirect(url_for('login'))

# 个人中心页面
@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    user = User.query.get(session['user_id'])
    permissions = PERMISSIONS[user.role]['permissions']
    
    # 准备统计数据
    if user.role == 'student':
        stats = {
            'courses': len(user.enrolled_courses),
            'assignments': user.completed_assignments,
            'messages': Message.query.filter_by(receiver_id=user.id, is_read=False).count()
        }
    elif user.role == 'teacher':
        stats = {
            'courses': len(user.courses_taught),
            'assignments': user.graded_assignments,
            'messages': Message.query.filter_by(receiver_id=user.id, is_read=False).count()
        }
    else:  # ai-assistant
        stats = {
            'courses': user.courses,
            'assignments': user.questions_answered,
            'messages': Message.query.filter_by(receiver_id=user.id, is_read=False).count()
        }
    
    return render_template('dashboard.html', user=user, permissions=permissions, stats=stats)

# 智能教案整理页面
@app.route('/text2ppt')
def run_text2ppt():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以使用此功能')
        return redirect(url_for('dashboard'))
    return render_template('text2ppt.html', user=user)

# 生成PPT接口
@app.route('/generate_ppt', methods=['POST'])
def generate_ppt():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'})
    
    content = request.form.get('content', '')
    theme = request.form.get('theme', '商务蓝 (Business Blue)')
    
    if not content:
        return jsonify({'success': False, 'message': '请输入PPT内容'})
    
    try:
        # 动态导入text2ppt模块中的类
        import text2ppt
        
        # 分析文本
        analyzer = text2ppt.SimpleTextAnalyzer(content)
        slides_data = analyzer.analyze()
        
        if not slides_data:
            return jsonify({'success': False, 'message': '无法解析文本内容'})
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(suffix='.pptx', delete=False)
        file_path = temp_file.name
        temp_file.close()
        
        # 初始化生成器并生成PPT
        generator = text2ppt.PPTGenerator(file_path, theme)
        
        for data in slides_data:
            if data["type"] == "title":
                generator.add_title_slide(data["title"], data["subtitle"])
            elif data["type"] == "content":
                generator.add_content_slide(data["title"], data["points"], data.get("image_keyword"))
        
        generator.save()
        
        # 发送文件给用户下载
        return send_file(file_path, as_attachment=True, download_name=f'presentation_{datetime.now().strftime("%Y%m%d%H%M%S")}.pptx', mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation')
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'生成失败：{str(e)}'})

# 学习课程页面
@app.route('/courses')
def courses():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取用户已选课程
    enrolled_courses = [sc.course for sc in user.enrolled_courses]
    
    return render_template('students_html/courses.html', user=user, enrolled_courses=enrolled_courses)

# 课程库页面
@app.route('/course_library')
def course_library():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取所有课程
    all_courses = Course.query.all()
    
    # 获取用户已选课程ID列表，用于前端显示是否已选课
    enrolled_course_ids = [sc.course_id for sc in user.enrolled_courses]
    
    return render_template('students_html/course_library.html', user=user, courses=all_courses, enrolled_course_ids=enrolled_course_ids)

# 选课功能
@app.route('/enroll_course/<int:course_id>', methods=['POST'])
def enroll_course(course_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 检查课程是否存在
    course = Course.query.get(course_id)
    if not course:
        flash('课程不存在')
        return redirect(url_for('course_library'))
    
    # 检查是否已经选课
    existing_enrollment = StudentCourse.query.filter_by(student_id=user.id, course_id=course_id).first()
    if existing_enrollment:
        flash('您已经选过这门课程')
        return redirect(url_for('course_library'))
    
    # 创建选课记录
    new_enrollment = StudentCourse(student_id=user.id, course_id=course_id)
    db.session.add(new_enrollment)
    db.session.commit()
    
    flash('选课成功')
    return redirect(url_for('course_library'))

# 退课功能
@app.route('/drop_course/<int:course_id>', methods=['POST'])
def drop_course(course_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 检查选课记录是否存在
    enrollment = StudentCourse.query.filter_by(student_id=user.id, course_id=course_id).first()
    if not enrollment:
        flash('您未选过这门课程')
        return redirect(url_for('courses'))
    
    # 删除选课记录
    db.session.delete(enrollment)
    db.session.commit()
    
    flash('退课成功')
    return redirect(url_for('courses'))

# 观看课程视频页面
@app.route('/watch_course/<int:course_id>')
def watch_course(course_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 学生需要检查是否已选该课程，教师则不需要
    if user.role == 'student':
        # 检查是否已选该课程
        enrollment = StudentCourse.query.filter_by(student_id=user.id, course_id=course_id).first()
        if not enrollment:
            flash('您未选过这门课程')
            return redirect(url_for('courses'))
    
    # 获取课程信息
    course = Course.query.get(course_id)
    if not course:
        flash('课程不存在')
        return redirect(url_for('courses'))
    
    return render_template('students_html/watch_course.html', user=user, course=course)

# 提交作业页面
@app.route('/assignments')
def assignments():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    return render_template('students_html/assignments.html', user=user)

# AI智能分析报告上传路由
@app.route('/upload_report', methods=['POST'])
def upload_report():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 检查是否有文件上传
    if 'report_file' not in request.files:
        flash('未检测到上传文件')
        return redirect(url_for('assignments'))
    
    report_file = request.files['report_file']
    
    # 检查文件是否为空
    if report_file.filename == '':
        flash('未选择上传文件')
        return redirect(url_for('assignments'))
    
    # 导入AI分析模块
    from AI_analysis import file_upload, ai_grading, report_comparison
    
    # 处理文件上传
    upload_result = file_upload.handle_file_upload(report_file)
    
    if not upload_result['success']:
        flash(upload_result['error'])
        return redirect(url_for('assignments'))
    
    # 读取文件内容
    file_content = file_upload.read_file_content(upload_result['file_path'])
    
    # 进行AI批改评分
    grading_result = ai_grading.grade_report(file_content)
    
    # 进行报告对比分析
    comparison_result = report_comparison.compare_reports(file_content)
    
    # 渲染分析结果页面
    return render_template(
        'students_html/ai_analysis_result.html',
        total_score=grading_result['total_score'],
        scores=grading_result['scores'],
        suggestions=grading_result['suggestions'],
        comparison=comparison_result
    )

# 查看成绩页面
@app.route('/grades')
def grades():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    return render_template('students_html/grades.html', user=user)



# 参与讨论页面
@app.route('/discussions')
def discussions():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    return render_template('discussions.html', user=user)

# 消息列表页面
@app.route('/messages')
def messages():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取用户收到的所有消息，按时间倒序排列
    messages = Message.query.filter_by(receiver_id=user.id).order_by(Message.created_at.desc()).all()
    
    # 标记所有未读消息为已读
    for msg in messages:
        if not msg.is_read:
            msg.is_read = True
    db.session.commit()
    
    return render_template('messages.html', user=user, messages=messages)

# 消息详情页面
@app.route('/message/<int:message_id>')
def message_detail(message_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取消息详情
    message = Message.query.get(message_id)
    if not message or message.receiver_id != user.id:
        flash('消息不存在或无权限查看')
        return redirect(url_for('messages'))
    
    # 标记消息为已读
    if not message.is_read:
        message.is_read = True
        db.session.commit()
    
    return render_template('message_detail.html', user=user, message=message)

# 删除消息
@app.route('/delete_message/<int:message_id>', methods=['POST'])
def delete_message(message_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取消息
    message = Message.query.get(message_id)
    if not message or message.receiver_id != user.id:
        flash('消息不存在或无权限操作')
        return redirect(url_for('messages'))

# ------------------------ AI助教功能 ------------------------ #

# AI助教页面
@app.route('/ai_tutoring')
def ai_tutoring():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取用户的AI会话历史
    sessions = AITutoringSession.query.filter_by(user_id=user.id).order_by(AITutoringSession.created_at.desc()).all()
    
    return render_template('ai_tutoring.html', user=user, sessions=sessions)

# 创建新的AI辅导会话
@app.route('/create_ai_session', methods=['POST'])
def create_ai_session():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    session_type = request.form['session_type']
    topic = request.form['topic']
    course_id = request.form.get('course_id')
    
    # 创建新会话
    new_session = AITutoringSession(
        user_id=user.id,
        course_id=course_id if course_id else None,
        session_type=session_type,
        topic=topic
    )
    db.session.add(new_session)
    db.session.commit()
    
    return redirect(url_for('ai_session_detail', session_id=new_session.id))

# AI辅导会话详情
@app.route('/ai_session/<int:session_id>')
def ai_session_detail(session_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取会话详情
    session = AITutoringSession.query.get(session_id)
    if not session or session.user_id != user.id:
        flash('会话不存在或无权限查看')
        return redirect(url_for('ai_tutoring'))
    
    # 获取会话消息
    messages = AITutoringMessage.query.filter_by(session_id=session_id).order_by(AITutoringMessage.created_at).all()
    
    return render_template('ai_session_detail.html', user=user, session=session, messages=messages)

# 发送AI辅导消息
@app.route('/send_ai_message/<int:session_id>', methods=['POST'])
def send_ai_message(session_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    content = request.form['content']
    
    # 获取会话
    session = AITutoringSession.query.get(session_id)
    if not session or session.user_id != user.id:
        flash('会话不存在或无权限操作')
        return redirect(url_for('ai_tutoring'))
    
    # 添加用户消息
    user_message = AITutoringMessage(
        session_id=session_id,
        user_id=user.id,
        sender='user',
        content=content
    )
    db.session.add(user_message)
    db.session.commit()
    
    # AI回复（模拟）
    ai_response = simulate_ai_response(content, session)
    
    ai_message = AITutoringMessage(
        session_id=session_id,
        user_id=user.id,  # 或者AI的user ID
        sender='ai',
        content=ai_response
    )
    db.session.add(ai_message)
    db.session.commit()
    
    return redirect(url_for('ai_session_detail', session_id=session_id))

# 模拟AI回复的函数
def simulate_ai_response(user_content, session):
    # 这里可以集成真实的AI API，如ChatGPT、通义千问等
    # 目前使用简单的模拟回复
    responses = [
        "这是一个很好的问题！让我来帮你解答...",
        "关于这个知识点，我可以给你一些建议...",
        "这个问题涉及到多个方面，让我详细解释一下...",
        "我理解你的疑问，以下是我的解答...",
        "根据你的问题，我认为可以从以下几个角度思考..."
    ]
    return f"[AI助教] {random.choice(responses)}\n\n{user_content}"  # 简单模拟，实际应调用AI API

# ------------------------ 项目组功能 ------------------------ #

# 项目组列表页面
@app.route('/project_groups')
def project_groups():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取用户参与的所有项目组
    user_groups = []
    for member in user.project_groups:
        user_groups.append(member.group)
    
    # 获取所有项目组（用于加入）
    all_groups = ProjectGroup.query.all()
    
    return render_template('project_groups.html', user=user, user_groups=user_groups, all_groups=all_groups)

# 创建项目组页面
@app.route('/create_project_group', methods=['GET', 'POST'])
def create_project_group():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    if request.method == 'POST':
        name = request.form['name']
        description = request.form['description']
        
        # 创建项目组
        new_group = ProjectGroup(
            name=name,
            description=description,
            creator_id=user.id
        )
        db.session.add(new_group)
        db.session.commit()
        
        # 添加创建者为管理员
        group_member = GroupMember(
            group_id=new_group.id,
            user_id=user.id,
            role='admin'
        )
        db.session.add(group_member)
        db.session.commit()
        
        flash('项目组创建成功')
        return redirect(url_for('project_groups'))
    
    return render_template('create_project_group.html', user=user)

# 项目组详情页面
@app.route('/project_group/<int:group_id>')
def project_group_detail(group_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取项目组详情
    group = ProjectGroup.query.get(group_id)
    if not group:
        flash('项目组不存在')
        return redirect(url_for('project_groups'))
    
    # 检查用户是否是成员
    is_member = any(member.user_id == user.id for member in group.members)
    
    # 获取项目组讨论消息
    discussions = GroupDiscussion.query.filter_by(group_id=group_id).order_by(GroupDiscussion.created_at).all()
    
    return render_template('project_group_detail.html', user=user, group=group, is_member=is_member, discussions=discussions)

# 加入项目组
@app.route('/join_project_group/<int:group_id>', methods=['POST'])
def join_project_group(group_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取项目组
    group = ProjectGroup.query.get(group_id)
    if not group:
        flash('项目组不存在')
        return redirect(url_for('project_groups'))
    
    # 检查是否已经是成员
    existing_member = GroupMember.query.filter_by(group_id=group_id, user_id=user.id).first()
    if existing_member:
        flash('您已经是该项目组的成员')
        return redirect(url_for('project_group_detail', group_id=group_id))
    
    # 添加成员
    new_member = GroupMember(
        group_id=group_id,
        user_id=user.id,
        role='member'
    )
    db.session.add(new_member)
    db.session.commit()
    
    flash('加入项目组成功')
    return redirect(url_for('project_group_detail', group_id=group_id))

# 退出项目组
@app.route('/leave_project_group/<int:group_id>', methods=['POST'])
def leave_project_group(group_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 获取成员记录
    member = GroupMember.query.filter_by(group_id=group_id, user_id=user.id).first()
    if not member:
        flash('您不是该项目组的成员')
        return redirect(url_for('project_groups'))
    
    # 删除成员记录
    db.session.delete(member)
    db.session.commit()
    
    flash('退出项目组成功')
    return redirect(url_for('project_groups'))

# 发送项目组讨论消息
@app.route('/send_group_message/<int:group_id>', methods=['POST'])
def send_group_message(group_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    content = request.form['content']
    
    # 检查用户是否是成员
    member = GroupMember.query.filter_by(group_id=group_id, user_id=user.id).first()
    if not member:
        flash('您不是该项目组的成员，无法发送消息')
        return redirect(url_for('project_group_detail', group_id=group_id))
    
    # 添加讨论消息
    new_message = GroupDiscussion(
        group_id=group_id,
        user_id=user.id,
        content=content
    )
    db.session.add(new_message)
    db.session.commit()
    
    return redirect(url_for('project_group_detail', group_id=group_id))

# 添加老师到项目组
@app.route('/add_teacher_to_group/<int:group_id>', methods=['POST'])
def add_teacher_to_group(group_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    
    # 检查用户是否是项目组管理员
    member = GroupMember.query.filter_by(group_id=group_id, user_id=user.id, role='admin').first()
    if not member:
        flash('只有管理员可以添加老师')
        return redirect(url_for('project_group_detail', group_id=group_id))
    
    teacher_username = request.form['teacher_username']
    
    # 查找老师
    teacher = User.query.filter_by(username=teacher_username, role='teacher').first()
    if not teacher:
        flash('老师不存在')
        return redirect(url_for('project_group_detail', group_id=group_id))
    
    # 检查老师是否已经是成员
    existing_member = GroupMember.query.filter_by(group_id=group_id, user_id=teacher.id).first()
    if existing_member:
        flash('老师已经是该项目组的成员')
        return redirect(url_for('project_group_detail', group_id=group_id))
    
    # 添加老师为成员
    new_member = GroupMember(
        group_id=group_id,
        user_id=teacher.id,
        role='teacher'
    )
    db.session.add(new_member)
    db.session.commit()
    
    flash('老师添加成功')
    return redirect(url_for('project_group_detail', group_id=group_id))
    
    # 删除消息
    db.session.delete(message)
    db.session.commit()
    
    flash('消息已删除')
    return redirect(url_for('messages'))

# 教师端-AI智能测验管理页面
@app.route('/teacher_ai_test_management')
def teacher_ai_test_management():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取教师创建的所有测验
    quizzes = Quiz.query.filter_by(teacher_id=user.id).all()
    
    return render_template('teacher_html/teacher_ai_test_management.html', user=user, quizzes=quizzes)

# 教师端-创建测验页面
@app.route('/teacher_create_quiz', methods=['GET', 'POST'])
def teacher_create_quiz():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取所有课程供教师选择
    courses = Course.query.all()
    
    if request.method == 'POST':
        # 处理测验创建表单提交，添加默认值处理
        title = request.form.get('title', f'测验_{int(time.time())}')
        course_id = int(request.form.get('course_id', 1))
        # 获取选中的知识点列表
        knowledge_points_list = request.form.getlist('knowledge_points')
        # 如果没有选中任何知识点，使用默认值
        if not knowledge_points_list:
            knowledge_points_list = ['设计基础']
        # 将列表转换为逗号分隔的字符串
        knowledge_points = ','.join(knowledge_points_list)
        difficulty = request.form.get('difficulty', '中等')
        time_limit = int(request.form.get('time_limit', 30))
        anti_cheat = 'anti_cheat' in request.form
        
        # 获取各种题型数量，确保至少为0
        single_count = max(0, int(request.form.get('single_count', 2)))
        multiple_count = max(0, int(request.form.get('multiple_count', 1)))
        judge_count = max(0, int(request.form.get('judge_count', 1)))
        short_count = max(0, int(request.form.get('short_count', 1)))
        discussion_count = max(0, int(request.form.get('discussion_count', 0)))
        
        # 生成唯一的测验ID
        quiz_id = f'quiz_{int(time.time())}'
        
        # 创建新测验
        new_quiz = Quiz(
            quiz_id=quiz_id,
            title=title,
            teacher_id=user.id,
            course_id=course_id,
            knowledge_points=knowledge_points,
            difficulty=difficulty,
            time_limit=time_limit,
            anti_cheat=anti_cheat
        )
        
        db.session.add(new_quiz)
        db.session.commit()
        
        # 从题库中抽取题目并添加到测验
        # 按题型和难度从题库中随机抽取题目
        
        # 抽取单选题
        if single_count > 0:
            # 查询符合条件的单选题，先尝试使用知识点过滤
            single_questions = Question.query.filter(
                Question.qtype == 'single',
                Question.difficulty == difficulty
            )
            # 先按知识点过滤
            if knowledge_points_list:
                single_questions = single_questions.filter(Question.knowledge_point.in_(knowledge_points_list))
            single_questions = single_questions.order_by(db.func.random()).limit(single_count).all()
            
            # 如果按知识点过滤没有找到足够的题目，尝试不使用知识点过滤
            if len(single_questions) < single_count:
                single_questions = Question.query.filter(
                    Question.qtype == 'single',
                    Question.difficulty == difficulty
                ).order_by(db.func.random()).limit(single_count).all()
            
            # 添加到测验
            for question in single_questions:
                quiz_question = QuizQuestion(
                    quiz_id=new_quiz.id,
                    question_id=question.id,
                    score=10.0  # 默认分值，可以根据需要调整
                )
                db.session.add(quiz_question)
        
        # 抽取多选题
        if multiple_count > 0:
            # 查询符合条件的多选题，先尝试使用知识点过滤
            multiple_questions = Question.query.filter(
                Question.qtype == 'multiple',
                Question.difficulty == difficulty
            )
            # 先按知识点过滤
            if knowledge_points_list:
                multiple_questions = multiple_questions.filter(Question.knowledge_point.in_(knowledge_points_list))
            multiple_questions = multiple_questions.order_by(db.func.random()).limit(multiple_count).all()
            
            # 如果按知识点过滤没有找到足够的题目，尝试不使用知识点过滤
            if len(multiple_questions) < multiple_count:
                multiple_questions = Question.query.filter(
                    Question.qtype == 'multiple',
                    Question.difficulty == difficulty
                ).order_by(db.func.random()).limit(multiple_count).all()
            
            for question in multiple_questions:
                quiz_question = QuizQuestion(
                    quiz_id=new_quiz.id,
                    question_id=question.id,
                    score=15.0
                )
                db.session.add(quiz_question)
        
        # 抽取判断题
        if judge_count > 0:
            # 查询符合条件的判断题，先尝试使用知识点过滤
            judge_questions = Question.query.filter(
                Question.qtype == 'judge',
                Question.difficulty == difficulty
            )
            # 先按知识点过滤
            if knowledge_points_list:
                judge_questions = judge_questions.filter(Question.knowledge_point.in_(knowledge_points_list))
            judge_questions = judge_questions.order_by(db.func.random()).limit(judge_count).all()
            
            # 如果按知识点过滤没有找到足够的题目，尝试不使用知识点过滤
            if len(judge_questions) < judge_count:
                judge_questions = Question.query.filter(
                    Question.qtype == 'judge',
                    Question.difficulty == difficulty
                ).order_by(db.func.random()).limit(judge_count).all()
            
            for question in judge_questions:
                quiz_question = QuizQuestion(
                    quiz_id=new_quiz.id,
                    question_id=question.id,
                    score=5.0
                )
                db.session.add(quiz_question)
        
        # 抽取简答题
        if short_count > 0:
            # 查询符合条件的简答题，先尝试使用知识点过滤
            short_questions = Question.query.filter(
                Question.qtype == 'short',
                Question.difficulty == difficulty
            )
            # 先按知识点过滤
            if knowledge_points_list:
                short_questions = short_questions.filter(Question.knowledge_point.in_(knowledge_points_list))
            short_questions = short_questions.order_by(db.func.random()).limit(short_count).all()
            
            # 如果按知识点过滤没有找到足够的题目，尝试不使用知识点过滤
            if len(short_questions) < short_count:
                short_questions = Question.query.filter(
                    Question.qtype == 'short',
                    Question.difficulty == difficulty
                ).order_by(db.func.random()).limit(short_count).all()
            
            for question in short_questions:
                quiz_question = QuizQuestion(
                    quiz_id=new_quiz.id,
                    question_id=question.id,
                    score=20.0
                )
                db.session.add(quiz_question)
        
        # 抽取论述题
        if discussion_count > 0:
            # 查询符合条件的论述题，先尝试使用知识点过滤
            discussion_questions = Question.query.filter(
                Question.qtype == 'short',  # 系统中论述题也存储为short类型
                Question.difficulty == difficulty
            )
            # 先按知识点过滤
            if knowledge_points_list:
                discussion_questions = discussion_questions.filter(Question.knowledge_point.in_(knowledge_points_list))
            discussion_questions = discussion_questions.order_by(db.func.random()).limit(discussion_count).all()
            
            # 如果按知识点过滤没有找到足够的题目，尝试不使用知识点过滤
            if len(discussion_questions) < discussion_count:
                discussion_questions = Question.query.filter(
                    Question.qtype == 'short',
                    Question.difficulty == difficulty
                ).order_by(db.func.random()).limit(discussion_count).all()
            
            for question in discussion_questions:
                quiz_question = QuizQuestion(
                    quiz_id=new_quiz.id,
                    question_id=question.id,
                    score=25.0
                )
                db.session.add(quiz_question)
        
        db.session.commit()
        
        flash('测验创建成功，已从题库中抽取题目')
        return redirect(url_for('teacher_ai_test_management'))
    
    return render_template('teacher_html/teacher_create_quiz.html', user=user, courses=courses)

# 教师端-查看测验详情页面
@app.route('/teacher_view_quiz/<int:quiz_id>')
def teacher_view_quiz(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取测验信息
    quiz = Quiz.query.get(quiz_id)
    if not quiz:
        flash('测验不存在')
        return redirect(url_for('teacher_ai_test_management'))
    
    # 获取测验的题目并预处理选项
    quiz_questions = QuizQuestion.query.filter_by(quiz_id=quiz_id).all()
    
    # 预处理每个题目的选项，解析JSON格式
    for quiz_q in quiz_questions:
        if quiz_q.question and quiz_q.question.options:
            try:
                # 解析JSON选项
                if quiz_q.question.options != 'null':
                    quiz_q.question.options_dict = json.loads(quiz_q.question.options)
                else:
                    quiz_q.question.options_dict = {}
            except (json.JSONDecodeError, TypeError):
                # 解析失败时设为空字典
                quiz_q.question.options_dict = {}
        else:
            quiz_q.question.options_dict = {}
    
    # 获取学生答题情况
    student_quizzes = StudentQuiz.query.filter_by(quiz_id=quiz_id).all()
    
    return render_template('teacher_html/teacher_view_quiz.html', user=user, quiz=quiz, quiz_questions=quiz_questions, student_quizzes=student_quizzes)

# 教师端-删除测验功能
@app.route('/teacher_delete_quiz/<int:quiz_id>', methods=['POST'])
def teacher_delete_quiz(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取测验信息
    quiz = Quiz.query.get(quiz_id)
    if not quiz:
        flash('测验不存在')
        return redirect(url_for('teacher_ai_test_management'))
    
    try:
        # 级联删除相关记录
        # 1. 删除学生答题记录（通过StudentQuiz关联）
        student_quizzes = StudentQuiz.query.filter_by(quiz_id=quiz_id).all()
        for student_quiz in student_quizzes:
            # 删除该学生测验的所有答题记录
            StudentAnswer.query.filter_by(student_quiz_id=student_quiz.id).delete()
            # 删除学生测验记录
            db.session.delete(student_quiz)
        
        # 2. 删除测验题目关联
        QuizQuestion.query.filter_by(quiz_id=quiz_id).delete()
        
        # 3. 删除测验本身
        db.session.delete(quiz)
        
        # 提交事务
        db.session.commit()
        
        flash('测验删除成功')
    except Exception as e:
        db.session.rollback()
        flash(f'测验删除失败: {str(e)}')
    
    return redirect(url_for('teacher_ai_test_management'))

# 教师端-上传题目页面
@app.route('/teacher_upload_question', methods=['GET', 'POST'])
def teacher_upload_question():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        # 检查是否是文件上传
        if 'question_file' in request.files and request.files['question_file'].filename != '':
            # 批量上传处理
            import csv
            from io import TextIOWrapper
            
            file = request.files['question_file']
            
            # 统计成功和失败数量
            success_count = 0
            error_count = 0
            
            # 解析题目数据的通用函数
            def parse_question_data(data_row):
                try:
                    # 检查行数据是否完整
                    if len(data_row) < 9:
                        return False
                    
                    # 解析行数据
                    knowledge_point = data_row[0].strip()
                    difficulty = data_row[1].strip()
                    qtype = data_row[2].strip()
                    content = data_row[3].strip()
                    option_a = data_row[4].strip()
                    option_b = data_row[5].strip()
                    option_c = data_row[6].strip()
                    option_d = data_row[7].strip()
                    answer = data_row[8].strip()
                    
                    # 验证必填字段
                    if not all([knowledge_point, difficulty, qtype, content, answer]):
                        return False
                    
                    # 验证题型
                    if qtype not in ['single', 'multiple', 'judge', 'short']:
                        return False
                    
                    # 验证难度
                    if difficulty not in ['简单', '中等', '困难']:
                        return False
                    
                    # 处理选项
                    options = None
                    if qtype in ['single', 'multiple', 'judge']:
                        options_dict = {}
                        if option_a:
                            options_dict['A'] = option_a
                        if option_b:
                            options_dict['B'] = option_b
                        if option_c:
                            options_dict['C'] = option_c
                        if option_d:
                            options_dict['D'] = option_d
                        if options_dict:
                            options = json.dumps(options_dict)
                    
                    # 生成唯一的题目ID，使用更精确的时间戳和随机数组合
                    qid = f'question_{int(time.time() * 1000)}_{random.randint(10000, 99999)}'
                    
                    # 创建新题目
                    new_question = Question(
                        qid=qid,
                        knowledge_point=knowledge_point,
                        difficulty=difficulty,
                        qtype=qtype,
                        content=content,
                        answer=answer,
                        options=options,
                        source='manual_upload'  # 手动上传标识
                    )
                    
                    db.session.add(new_question)
                    return True
                except Exception as e:
                    return False
            
            if file.filename.endswith('.csv'):
                # 读取CSV文件
                csv_file = TextIOWrapper(file, encoding='utf-8')
                reader = csv.reader(csv_file)
                
                # 跳过表头
                next(reader, None)
                
                for row in reader:
                    if parse_question_data(row):
                        success_count += 1
                    else:
                        error_count += 1
            elif file.filename.endswith('.docx'):
                # 读取DOCX文件
                from docx import Document
                
                doc = Document(file)
                # 提取所有段落文本
                all_text = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_text.append(text)
                
                # 跳过表头（如果有）
                if all_text:
                    # 检查第一行是否为表头，如果是则跳过
                    first_line = all_text[0]
                    if '知识点' in first_line and '难度' in first_line and '题型' in first_line:
                        all_text = all_text[1:]
                    
                    # 处理每一行
                    for line in all_text:
                        # 按逗号分隔行数据
                        row = line.split(',')
                        if parse_question_data(row):
                            success_count += 1
                        else:
                            error_count += 1
            elif file.filename.endswith('.html'):
                # 读取HTML文件
                from parse_html_questions import parse_html_questions
                
                html_content = file.read().decode('utf-8')
                # 解析HTML内容
                questions_data = parse_html_questions(html_content)
                
                # 处理解析得到的题目数据
                for question_data in questions_data:
                    try:
                        # 生成唯一的题目ID
                        qid = f'question_{int(time.time())}_{random.randint(1000, 9999)}'
                        
                        # 创建新题目
                        new_question = Question(
                            qid=qid,
                            knowledge_point=question_data['knowledge_point'],
                            difficulty=question_data['difficulty'],
                            qtype=question_data['qtype'],
                            content=question_data['question_text'],
                            answer=question_data['answer'],
                            options=question_data['options'],
                            source=question_data['source']
                        )
                        
                        db.session.add(new_question)
                        success_count += 1
                    except Exception as e:
                        error_count += 1
            else:
                flash('请上传CSV、DOCX或HTML格式的文件')
                return redirect(url_for('teacher_upload_question'))
            
            # 提交数据库更改
            db.session.commit()
            
            flash(f'批量上传完成！成功上传 {success_count} 道题目，失败 {error_count} 道题目')
            return redirect(url_for('teacher_upload_question'))
        
        # 单个题目上传处理
        # 获取表单数据
        knowledge_point = request.form['knowledge_point']
        difficulty = request.form['difficulty']
        qtype = request.form['qtype']
        content = request.form['content']
        answer = request.form['answer']
        
        # 处理选项
        options = None
        if qtype in ['single', 'multiple', 'judge']:
            options_dict = {}
            for i in ['A', 'B', 'C', 'D']:
                option_value = request.form.get(f'option_{i}')
                if option_value:
                    options_dict[i] = option_value
            if options_dict:
                options = json.dumps(options_dict)
        
        # 生成唯一的题目ID
        qid = f'question_{int(time.time())}_{random.randint(1000, 9999)}'
        
        # 创建新题目
        new_question = Question(
            qid=qid,
            knowledge_point=knowledge_point,
            difficulty=difficulty,
            qtype=qtype,
            content=content,
            answer=answer,
            options=options,
            source='manual_upload'  # 手动上传标识
        )
        
        db.session.add(new_question)
        db.session.commit()
        
        flash('题目上传成功')
        return redirect(url_for('teacher_upload_question'))
    
    return render_template('teacher_html/teacher_upload_question.html', user=user)

# API端点：直接上传题库文件到后台
@app.route('/api/upload_question', methods=['POST'])
def api_upload_question():
    """
    直接上传题库文件到后台的API端点
    支持CSV和DOCX格式
    无需登录验证，直接上传到数据库
    """
    try:
        if 'question_file' not in request.files:
            return jsonify({'success': False, 'message': '未找到文件'}), 400
        
        file = request.files['question_file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '未选择文件'}), 400
        
        # 检查文件格式
        if not (file.filename.endswith('.csv') or file.filename.endswith('.docx') or file.filename.endswith('.html')):
            return jsonify({'success': False, 'message': '请上传CSV、DOCX或HTML格式的文件'}), 400
        
        # 统计成功和失败数量
        success_count = 0
        error_count = 0
        
        # 解析题目数据的通用函数
        def parse_question_data(data_row):
            try:
                # 检查行数据是否完整
                if len(data_row) < 9:
                    return False
                
                # 解析行数据
                knowledge_point = data_row[0].strip()
                difficulty = data_row[1].strip()
                qtype = data_row[2].strip()
                content = data_row[3].strip()
                option_a = data_row[4].strip()
                option_b = data_row[5].strip()
                option_c = data_row[6].strip()
                option_d = data_row[7].strip()
                answer = data_row[8].strip()
                
                # 验证必填字段
                if not all([knowledge_point, difficulty, qtype, content, answer]):
                    return False
                
                # 验证题型
                if qtype not in ['single', 'multiple', 'judge', 'short']:
                    return False
                
                # 验证难度
                if difficulty not in ['简单', '中等', '困难']:
                    return False
                
                # 处理选项
                options = None
                if qtype in ['single', 'multiple', 'judge']:
                    options_dict = {}
                    if option_a:
                        options_dict['A'] = option_a
                    if option_b:
                        options_dict['B'] = option_b
                    if option_c:
                        options_dict['C'] = option_c
                    if option_d:
                        options_dict['D'] = option_d
                    if options_dict:
                        options = json.dumps(options_dict)
                
                # 生成唯一的题目ID，使用更精确的时间戳和随机数组合
                qid = f'question_{int(time.time() * 1000)}_{random.randint(10000, 99999)}'
                
                # 创建新题目
                new_question = Question(
                    qid=qid,
                    knowledge_point=knowledge_point,
                    difficulty=difficulty,
                    qtype=qtype,
                    content=content,
                    answer=answer,
                    options=options,
                    source='manual_upload'  # 手动上传标识
                )
                
                db.session.add(new_question)
                return True
            except Exception as e:
                return False
        
        if file.filename.endswith('.csv'):
            # 读取CSV文件
            import csv
            from io import TextIOWrapper
            csv_file = TextIOWrapper(file, encoding='utf-8')
            reader = csv.reader(csv_file)
            
            # 跳过表头
            next(reader, None)
            
            for row in reader:
                if parse_question_data(row):
                    success_count += 1
                else:
                    error_count += 1
        elif file.filename.endswith('.docx'):
            # 读取DOCX文件
            from docx import Document
            
            doc = Document(file)
            # 提取所有段落文本
            all_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text.append(text)
            
            # 跳过表头（如果有）
            if all_text:
                # 检查第一行是否为表头，如果是则跳过
                first_line = all_text[0]
                if '知识点' in first_line and '难度' in first_line and '题型' in first_line:
                    all_text = all_text[1:]
                
                # 处理每一行
                for line in all_text:
                    # 按逗号分隔行数据
                    row = line.split(',')
                    if parse_question_data(row):
                        success_count += 1
                    else:
                        error_count += 1
        elif file.filename.endswith('.html'):
            # 读取HTML文件
            from parse_html_questions import parse_html_questions
            
            html_content = file.read().decode('utf-8')
            # 解析HTML内容
            questions_data = parse_html_questions(html_content)
            
            # 处理解析得到的题目数据
            for question_data in questions_data:
                try:
                    # 生成唯一的题目ID
                    qid = f'question_{int(time.time() * 1000)}_{random.randint(10000, 99999)}'
                    
                    # 创建新题目
                    new_question = Question(
                        qid=qid,
                        knowledge_point=question_data['knowledge_point'],
                        difficulty=question_data['difficulty'],
                        qtype=question_data['qtype'],
                        content=question_data['question_text'],
                        answer=question_data['answer'],
                        options=question_data['options'],
                        source=question_data['source']
                    )
                    
                    db.session.add(new_question)
                    success_count += 1
                except Exception as e:
                    error_count += 1
        
        # 提交数据库更改
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '文件上传成功',
            'success_count': success_count,
            'error_count': error_count
        }), 200
    except Exception as e:
        return jsonify({'success': False, 'message': f'上传失败：{str(e)}'}), 500

# 教师端-查看班级学生答题情况
@app.route('/teacher_view_student_answers/<int:quiz_id>')
def teacher_view_student_answers(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取测验信息
    quiz = Quiz.query.get(quiz_id)
    if not quiz:
        flash('测验不存在')
        return redirect(url_for('teacher_ai_test_management'))
    
    # 获取所有学生的答题记录
    student_quizzes = StudentQuiz.query.filter_by(quiz_id=quiz_id).all()
    
    return render_template('teacher_html/teacher_view_student_answers.html', user=user, quiz=quiz, student_quizzes=student_quizzes)

# 教师端-查看错题库
@app.route('/teacher_view_error_bank')
def teacher_view_error_bank():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取所有错题库记录
    error_bank = ErrorQuestionBank.query.all()
    
    return render_template('teacher_html/teacher_view_error_bank.html', user=user, error_bank=error_bank)

# 学生端-AI智能测验列表页面
@app.route('/student_ai_quizzes')
def student_ai_quizzes():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取课程ID参数
    course_id = request.args.get('course_id')
    
    # 获取学生已选课程
    enrolled_courses = StudentCourse.query.filter_by(student_id=user.id).all()
    enrolled_course_ids = [ec.course_id for ec in enrolled_courses]
    
    # 根据课程ID过滤测验
    if course_id:
        course_id = int(course_id)
        # 确保学生已选该课程
        if course_id in enrolled_course_ids:
            quizzes = Quiz.query.filter_by(course_id=course_id).all()
        else:
            flash('您未选过该课程')
            quizzes = []
    else:
        # 获取学生已选课程的所有测验
        quizzes = Quiz.query.filter(Quiz.course_id.in_(enrolled_course_ids)).all()
    
    # 获取学生已参与的测验
    student_quizzes = StudentQuiz.query.filter_by(student_id=user.id).all()
    
    return render_template('students_html/student_ai_quizzes.html', user=user, quizzes=quizzes, student_quizzes=student_quizzes, course_id=course_id)

# 学生端-独立AI测验入口
@app.route('/student_self_quiz')
def student_self_quiz():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取学生已选课程
    enrolled_courses = [sc.course for sc in user.enrolled_courses]
    
    return render_template('students_html/student_self_quiz.html', user=user, courses=enrolled_courses)

# 学生端-生成独立AI测验
@app.route('/student_generate_self_quiz', methods=['POST'])
def student_generate_self_quiz():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取表单数据
    knowledge_points = request.form.get('knowledge_points', 'Python基础')
    difficulty = request.form.get('difficulty', '中等')
    single_count = int(request.form.get('single_count', 2))
    multiple_count = int(request.form.get('multiple_count', 1))
    judge_count = int(request.form.get('judge_count', 1))
    short_count = int(request.form.get('short_count', 1))
    time_limit = int(request.form.get('time_limit', 30))
    
    # 从AI_test_system模块导入功能
    from AI_test_system import AIQuestionGenerator, TestConfig
    
    # 创建测验配置
    test_config = TestConfig(
        test_id=f'self_{user.id}_{int(time.time())}',
        knowledge_points=knowledge_points.split(','),
        difficulty=difficulty,
        qtype_counts={
            'single': single_count,
            'multiple': multiple_count,
            'judge': judge_count,
            'short': short_count
        },
        time_limit=time_limit,
        anti_cheat=True
    )
    
    # 生成题目
    generator = AIQuestionGenerator()
    questions = generator.generate_ai_questions(test_config)
    
    # 保存题目到数据库
    for q in questions:
        # 检查题目是否已存在
        existing_question = Question.query.filter_by(qid=q.qid).first()
        if not existing_question:
            new_question = Question(
                qid=q.qid,
                knowledge_point=q.knowledge_point,
                difficulty=q.difficulty,
                qtype=q.qtype,
                content=q.content,
                answer=q.answer,
                options=json.dumps(q.options) if q.options else None,
                score_std=q.score_std,
                source=q.source
            )
            db.session.add(new_question)
            db.session.commit()
    
    db.session.commit()
    
    # 创建一个临时测验记录（用于学生答题）
    temp_quiz = Quiz(
        quiz_id=f'self_{user.id}_{int(time.time())}',
        title='自主AI测验',
        teacher_id=1,  # 默认教师ID
        course_id=1,  # 默认课程ID
        knowledge_points=knowledge_points,
        difficulty=difficulty,
        time_limit=time_limit,
        anti_cheat=True
    )
    db.session.add(temp_quiz)
    db.session.commit()
    
    # 添加题目到测验
    for q in questions:
        question = Question.query.filter_by(qid=q.qid).first()
        if question:
            quiz_question = QuizQuestion(
                quiz_id=temp_quiz.id,
                question_id=question.id,
                score=10.0  # 默认每题10分
            )
            db.session.add(quiz_question)
    
    db.session.commit()
    
    # 创建学生测验记录
    student_quiz = StudentQuiz(
        student_id=user.id,
        quiz_id=temp_quiz.id,
        start_time=datetime.utcnow(),
        status='in_progress'
    )
    db.session.add(student_quiz)
    db.session.commit()
    
    return render_template('students_html/student_start_quiz.html', user=user, quiz=temp_quiz, questions=questions, student_quiz_id=student_quiz.id)

# 学生端-开始测验页面
@app.route('/student_start_quiz/<int:quiz_id>')
def student_start_quiz(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取测验信息
    quiz = Quiz.query.get(quiz_id)
    if not quiz:
        flash('测验不存在')
        return redirect(url_for('student_ai_quizzes'))
    
    # 获取测验关联的题目
    quiz_questions = QuizQuestion.query.filter_by(quiz_id=quiz.id).all()
    if not quiz_questions:
        flash('该测验暂无题目')
        return redirect(url_for('student_ai_quizzes'))
    
    # 获取题目详情
    questions = []
    for qq in quiz_questions:
        question = Question.query.get(qq.question_id)
        if question:
            questions.append(question)
    
    # 预处理题目选项，不修改数据库对象
    for question in questions:
        if question.options:
            # 创建临时属性存储解析后的选项
            question.parsed_options = json.loads(question.options)
        else:
            question.parsed_options = {}
    
    # 检查是否已经存在进行中的学生测验记录
    existing_student_quiz = StudentQuiz.query.filter_by(
        student_id=user.id,
        quiz_id=quiz.id,
        status='in_progress'
    ).first()
    
    if existing_student_quiz:
        # 使用已存在的测验记录
        student_quiz = existing_student_quiz
    else:
        # 创建新的学生测验记录
        student_quiz = StudentQuiz(
            student_id=user.id,
            quiz_id=quiz.id,
            start_time=datetime.utcnow(),
            status='in_progress'
        )
        db.session.add(student_quiz)
        db.session.commit()
    
    return render_template('students_html/student_start_quiz.html', user=user, quiz=quiz, questions=questions, student_quiz_id=student_quiz.id)

# 学生端-提交答案
@app.route('/student_submit_answer', methods=['POST'])
def student_submit_answer():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取提交的数据
    student_quiz_id = int(request.form['student_quiz_id'])
    question_id = int(request.form['question_id'])
    student_answer = request.form['student_answer']
    spend_time = float(request.form['spend_time'])
    
    # 创建学生答题记录
    student_answer_record = StudentAnswer(
        student_quiz_id=student_quiz_id,
        question_id=question_id,
        student_answer=student_answer,
        spend_time=spend_time
    )
    
    db.session.add(student_answer_record)
    db.session.commit()
    
    return jsonify({'status': 'success'})

# 学生端-结束测验
@app.route('/student_end_quiz/<int:student_quiz_id>', methods=['GET', 'POST'])
def student_end_quiz(student_quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取学生测验记录
    student_quiz = StudentQuiz.query.get(student_quiz_id)
    if not student_quiz:
        flash('测验记录不存在')
        return redirect(url_for('student_ai_quizzes'))
    
    # 更新测验状态
    student_quiz.end_time = datetime.utcnow()
    student_quiz.status = 'completed'
    
    # 从AI_test_system模块导入功能
    from AI_test_system import IntelligentGrader, Question as AIQuestion
    
    # 获取学生所有答题记录
    student_answers = StudentAnswer.query.filter_by(student_quiz_id=student_quiz_id).all()
    
    # 获取对应的题目
    questions = []
    for sa in student_answers:
        question = Question.query.get(sa.question_id)
        if question:
            # 转换为AIQuestion对象
            options = json.loads(question.options) if question.options else None
            ai_question = AIQuestion(
                qid=question.qid,
                knowledge_point=question.knowledge_point,
                difficulty=question.difficulty,
                qtype=question.qtype,
                content=question.content,
                answer=question.answer,
                options=options,
                score_std=question.score_std
            )
            questions.append(ai_question)
    
    # 创建AnswerRecord对象列表
    from AI_test_system import AnswerRecord
    answer_records = []
    for sa in student_answers:
        answer_record = AnswerRecord(
            sid=str(user.id),
            qid=Question.query.get(sa.question_id).qid,
            test_id=student_quiz.quiz.quiz_id,
            student_answer=sa.student_answer,
            spend_time=sa.spend_time
        )
        answer_records.append(answer_record)
    
    # 智能批改
    grader = IntelligentGrader()
    graded_records = grader.batch_grade(answer_records, questions)
    
    # 更新学生答题记录
    total_score = 0.0
    for i, sa in enumerate(student_answers):
        sa.is_correct = graded_records[i].is_correct
        sa.score = graded_records[i].score
        total_score += graded_records[i].score
        
        # 如果答错，添加到错题库
        if not graded_records[i].is_correct:
            # 检查是否已存在于错题库
            existing_error = ErrorQuestionBank.query.filter_by(
                student_id=user.id,
                question_id=sa.question_id
            ).first()
            
            if existing_error:
                existing_error.error_count += 1
                existing_error.last_error_time = datetime.utcnow()
            else:
                new_error = ErrorQuestionBank(
                    student_id=user.id,
                    question_id=sa.question_id
                )
                db.session.add(new_error)
    
    # 更新测验总分
    student_quiz.total_score = total_score
    
    db.session.commit()
    
    return redirect(url_for('student_quiz_result', student_quiz_id=student_quiz_id))

# 学生端-测验结果页面
@app.route('/student_quiz_result/<int:student_quiz_id>')
def student_quiz_result(student_quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'student':
        flash('只有学生角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取学生测验记录
    student_quiz = StudentQuiz.query.get(student_quiz_id)
    if not student_quiz:
        flash('测验记录不存在')
        return redirect(url_for('student_ai_quizzes'))
    
    # 获取测验信息
    quiz = student_quiz.quiz
    
    # 获取学生答题记录
    student_answers = StudentAnswer.query.filter_by(student_quiz_id=student_quiz_id).all()
    
    # 计算总分数
    total_score = student_quiz.total_score
    
    # 计算答题用时
    time_spent = 0
    if student_quiz.start_time and student_quiz.end_time:
        time_spent = int((student_quiz.end_time - student_quiz.start_time).total_seconds() / 60)
    
    # 预处理题目选项，不修改数据库对象
    for sa in student_answers:
        if sa.question.options:
            # 创建临时属性存储解析后的选项
            sa.question.parsed_options = json.loads(sa.question.options)
        else:
            sa.question.parsed_options = {}
    
    return render_template('students_html/student_quiz_result.html', 
                         user=user, 
                         student_quiz=student_quiz, 
                         student_answers=student_answers,
                         quiz=quiz,
                         total_score=total_score,
                         time_spent=time_spent)

# 学生端-添加题目到错题集
@app.route('/add_to_error_bank', methods=['POST'])
def add_to_error_bank():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    student_id = session['user_id']
    question_id = request.form.get('question_id')
    student_quiz_id = request.form.get('student_quiz_id')
    
    if not question_id:
        flash('题目ID不存在')
        return redirect(url_for('student_quiz_result', student_quiz_id=student_quiz_id))
    
    # 检查是否已存在于错题库
    existing_error = ErrorQuestionBank.query.filter_by(
        student_id=student_id,
        question_id=question_id
    ).first()
    
    if existing_error:
        existing_error.error_count += 1
        existing_error.last_error_time = datetime.utcnow()
        flash('已添加到错题集（重复）')
    else:
        new_error = ErrorQuestionBank(
            student_id=student_id,
            question_id=question_id
        )
        db.session.add(new_error)
        flash('成功添加到错题集')
    
    db.session.commit()
    return redirect(url_for('student_quiz_result', student_quiz_id=student_quiz_id))

# 设置Flask应用的默认编码为UTF-8
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = 'application/json; charset=utf-8'

# 设置模板渲染的编码为UTF-8
app.jinja_env.auto_reload = True
app.jinja_env.encoding = 'utf-8'

# 添加全局响应拦截器，确保所有HTML响应都设置正确的UTF-8编码
@app.after_request
def set_utf8_content_type(response):
    if response.content_type.startswith('text/html'):
        response.headers['Content-Type'] = 'text/html; charset=utf-8'
    return response

# AI报告模块路由 - 暂时显示提示信息


# 确保上传目录存在
UPLOAD_FOLDER = 'course_files'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# 添加静态文件路由，用于访问课程视频
@app.route('/course_files/<path:filename>')
def serve_course_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

# 教师端-创建课程页面
@app.route('/teacher_create_course', methods=['GET', 'POST'])
def teacher_create_course():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取当前老师创建的所有课程
    created_courses = Course.query.filter_by(teacher_id=user.id).all()
    
    if request.method == 'POST':
        # 处理课程创建表单提交
        course_code = request.form['course_code']
        title = request.form['title']
        description = request.form['description']
        credit = float(request.form['credit'])
        
        # 检查课程代码是否已存在
        existing_course = Course.query.filter_by(course_code=course_code).first()
        if existing_course:
            flash('课程代码已存在')
            return redirect(url_for('teacher_create_course'))
        
        # 处理文件上传
        video_path = None
        if 'course_file' in request.files:
            course_file = request.files['course_file']
            if course_file.filename != '':
                # 保存文件到上传目录
                filename = f"{course_code}_{course_file.filename}"
                video_path = os.path.join(UPLOAD_FOLDER, filename)
                course_file.save(video_path)
        
        # 创建新课程
        new_course = Course(
            course_code=course_code,
            title=title,
            description=description,
            credit=credit,
            teacher_id=user.id,
            video_path=video_path
        )
        
        db.session.add(new_course)
        db.session.commit()
        
        flash('课程上传成功')
        return redirect(url_for('teacher_create_course'))
    
    return render_template('teacher_html/teacher_create_course.html', user=user, created_courses=created_courses)

# 教师端-编辑课程页面
@app.route('/teacher_edit_course')
def teacher_edit_course():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    return render_template('teacher_html/teacher_edit_course.html', user=user)

# 教师端-批改作业页面
@app.route('/teacher_grade_assignments')
def teacher_grade_assignments():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取教师授课的所有课程
    courses = user.courses_taught
    
    # 获取所有相关测验
    quizzes = Quiz.query.filter_by(teacher_id=user.id).all()
    
    # 获取需要批改的学生测验（已提交但未批改的测验）
    student_quizzes = StudentQuiz.query.join(Quiz).filter(
        Quiz.teacher_id == user.id,
        StudentQuiz.status == 'completed'  # 已完成的测验
    ).order_by(StudentQuiz.submit_time.desc()).all()
    
    # 准备数据：按课程分组的测验记录
    course_quiz_data = {}
    for course in courses:
        course_quizzes = Quiz.query.filter_by(course_id=course.id, teacher_id=user.id).all()
        course_quiz_data[course.id] = {
            'course': course,
            'quizzes': course_quizzes,
            'student_records': []
        }
    
    # 为每个课程添加学生测验记录
    for student_quiz in student_quizzes:
        course_id = student_quiz.quiz.course_id
        if course_id in course_quiz_data:
            course_quiz_data[course_id]['student_records'].append(student_quiz)
    
    return render_template('teacher_html/teacher_grade_assignments.html', 
                           user=user, 
                           courses=courses, 
                           course_quiz_data=course_quiz_data, 
                           student_quizzes=student_quizzes)

# 教师端-管理学生页面
@app.route('/teacher_manage_students')
def teacher_manage_students():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    return render_template('teacher_html/teacher_manage_students.html', user=user)

# 教师端-发布成绩页面
@app.route('/teacher_publish_grades')
def teacher_publish_grades():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    
    # 获取教师授课的所有课程
    courses = user.courses_taught
    
    # 获取所有相关测验
    quizzes = Quiz.query.filter_by(teacher_id=user.id).all()
    
    # 获取所有学生测验成绩
    student_quizzes = StudentQuiz.query.join(Quiz).filter(
        Quiz.teacher_id == user.id,
        StudentQuiz.status == 'completed'  # 已完成的测验
    ).order_by(StudentQuiz.submit_time.desc()).all()
    
    # 准备数据：按课程分组的成绩记录
    course_grade_data = {}
    for course in courses:
        # 获取该课程下的所有学生
        students = User.query.join(StudentCourse).filter(
            StudentCourse.course_id == course.id,
            User.role == 'student'
        ).all()
        
        # 获取该课程下的所有测验
        course_quizzes = Quiz.query.filter_by(course_id=course.id, teacher_id=user.id).all()
        
        # 为每个学生获取该课程下的测验成绩
        student_grades = []
        for student in students:
            student_course_quizzes = StudentQuiz.query.join(Quiz).filter(
                StudentQuiz.student_id == student.id,
                Quiz.course_id == course.id,
                Quiz.teacher_id == user.id
            ).all()
            
            if student_course_quizzes:
                student_grades.append({
                    'student': student,
                    'quizzes': student_course_quizzes
                })
        
        course_grade_data[course.id] = {
            'course': course,
            'quizzes': course_quizzes,
            'student_grades': student_grades
        }
    
    return render_template('teacher_html/teacher_publish_grades.html', 
                           user=user, 
                           courses=courses, 
                           course_grade_data=course_grade_data, 
                           student_quizzes=student_quizzes)



# 教师端-参与讨论页面
@app.route('/teacher_discussions')
def teacher_discussions():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    if user.role != 'teacher':
        flash('只有教师角色可以访问此页面')
        return redirect(url_for('dashboard'))
    return render_template('teacher_html/teacher_discussions.html', user=user)

# 登出
@app.route('/logout')
def logout():
    session.clear()
    flash('已退出登录')
    return redirect(url_for('login'))

if __name__ == '__main__':
    app.run(debug=True, port=5000)